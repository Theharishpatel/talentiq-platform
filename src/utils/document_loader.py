"""
Document Loader.

Loads a Job Description from supported
document formats.

Currently Supported

- .docx
- .txt
"""

from pathlib import Path

from docx import Document


def _load_docx(
    path: Path,
) -> str:
    """
    Load text from a DOCX file.
    """

    document = Document(path)

    lines = []

    # Paragraphs
    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            lines.append(text)

    # Tables (future-proof)
    for table in document.tables:

        for row in table.rows:

            values = []

            for cell in row.cells:

                value = cell.text.strip()

                if value:
                    values.append(value)

            if values:
                lines.append(" | ".join(values))

    return "\n".join(lines)


def _load_txt(
    path: Path,
) -> str:
    """
    Load text from TXT file.
    """

    return path.read_text(
        encoding="utf-8",
    )


def load_document(
    path: str | Path,
) -> str:
    """
    Load supported document formats.

    Parameters
    ----------
    path : str | Path

    Returns
    -------
    str
        Document text.
    """

    path = Path(path)

    if not path.exists():

        raise FileNotFoundError(
            f"{path} not found."
        )

    suffix = path.suffix.lower()

    if suffix == ".docx":

        return _load_docx(path)

    if suffix == ".txt":

        return _load_txt(path)

    raise ValueError(
        f"Unsupported document format: {suffix}"
    )